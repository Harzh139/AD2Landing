from docx import Document
from docx.shared import Pt
import os

def create_docx():
    doc = Document()
    
    # Title
    doc.add_heading('Ad2Page: AI Landing Page Personalizer', 0)
    
    # Brief Explanation
    doc.add_heading('1. The Vision: From Generic to Surgical', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    p.add_run('Most digital marketing spend is wasted because users click highly specific ads but land on generic, one-size-fits-all pages. This "Message Mismatch" kills conversion rates.')
    
    p = doc.add_paragraph()
    run = p.add_run('Solution: ')
    run.bold = True
    p.add_run('I built an AI-powered CRO engine that surgically modifies existing landing pages. Instead of building new pages from scratch, I use AI to bridge the gaps in copy, tone, and urgency based on the ad creative.')
    
    p = doc.add_paragraph()
    run = p.add_run('Outcome: ')
    run.bold = True
    p.add_run('Marketers get instant, hyper-personalized variations that maintain brand identity while maximizing ROI.')

    # System Flow
    doc.add_heading('2. System Flow & Agent Architecture', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    p.add_run('Traditional automation tools are too rigid to understand the nuance of a specific marketing campaign.')
    
    p = doc.add_paragraph()
    run = p.add_run('Solution: ')
    run.bold = True
    p.add_run('I designed a modular agentic pipeline:')
    
    agents = [
        ('Analysis Agent', 'I use Llama 3.2 Vision to deconstruct psychological triggers from ad images.'),
        ('Scraper Agent', 'I built a robust crawler to capture the "Grounded Truth" of the existing website.'),
        ('CRO Strategy Agent', 'I prompt the AI as a Senior Copywriter to identify incremental copy optimizations.'),
        ('Logic Engine', 'I use a deterministic template builder to inject AI variants into clean HTML.')
    ]
    for title, desc in agents:
        li = doc.add_paragraph(style='List Bullet')
        r = li.add_run(f'{title}: ')
        r.bold = True
        li.add_run(desc)
    
    p = doc.add_paragraph()
    run = p.add_run('Outcome: ')
    run.bold = True
    p.add_run('A seamless flow from Ad Click to Customized Experience in under 10 seconds.')

    # Reliability
    doc.add_heading('3. Reliability: Solving the "AI Chaos" Problem', level=1)

    # Hallucinations
    doc.add_heading('A. Handling Hallucinations', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    p.add_run('LLMs often invent product features or use non-existent assets when they start "thinking" like creators.')
    
    p = doc.add_paragraph()
    run = p.add_run('Solution: ')
    run.bold = True
    p.add_run('I implement strict Context Grounding. I feed the raw scraped data into the AI and punish it if it deviates from the "Ground Truth." I also force it to inherit only the specific media assets found during scraping.')
    
    p = doc.add_paragraph()
    run = p.add_run('Outcome: ')
    run.bold = True
    p.add_run('The generated page stays 100% authentic to the real brand.')

    # Broken UI
    doc.add_heading('B. Handling Broken UI', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    p.add_run('Direct AI-generated code is brittle and often breaks layouts on different screen sizes.')
    
    p = doc.add_paragraph()
    run = p.add_run('Solution: ')
    run.bold = True
    p.add_run('I decoupled content from layout. I never let the AI write CSS or core HTML structure. Instead, it outputs structured JSON, and my deterministic Template Builder handles the technical rendering.')
    
    p = doc.add_paragraph()
    run = p.add_run('Outcome: ')
    run.bold = True
    p.add_run('Production-grade pages that are fully responsive and bug-free.')

    # Consistancy
    doc.add_heading('C. Handling Inconsistent Outputs', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    p.add_run('AI responses can be unpredictable, resulting in different structures every time you run it.')
    
    p = doc.add_paragraph()
    run = p.add_run('Solution: ')
    run.bold = True
    p.add_run('I enforced strict Schema Governance. I use JSON-response-format prompts with hard-coded keys and lock the variation count to exactly three.')
    
    p = doc.add_paragraph()
    run = p.add_run('Outcome: ')
    run.bold = True
    p.add_run('Predictable, high-quality variations every single time.')

    # Save
    filename = 'Ad2Page_Documentation_v2.docx'
    doc.save(filename)
    print(f'Successfully created {os.path.abspath(filename)}')

if __name__ == '__main__':
    create_docx()
